"""
=========================================================
DOCX Processing using Python
Libraries:
1. python-docx
2. docx2txt
3. mammoth
=========================================================
"""

from docx import Document
import docx2txt
import mammoth
import pandas as pd
from PIL import Image
import matplotlib.pyplot as plt
import io
import os

FILE = "file-sample_100kB.docx"

# =========================================================
# 1. python-docx
# =========================================================

print("=" * 70)
print("1. PYTHON-DOCX")
print("=" * 70)

doc = Document(FILE)

# ---------------------------------------------------------
# Paragraphs
# ---------------------------------------------------------

print("\nTotal Paragraphs:", len(doc.paragraphs))

print("\nReading Paragraphs\n")

for i, para in enumerate(doc.paragraphs, start=1):

    if para.text.strip():

        print(f"Paragraph {i}")
        print(para.text)
        print("-" * 50)



# ---------------------------------------------------------
# Metadata
# ---------------------------------------------------------

print("\nMetadata\n")

meta = doc.core_properties

print("Author   :", meta.author)

print("Title    :", meta.title)

print("Subject  :", meta.subject)

print("Created  :", meta.created)

print("Modified :", meta.modified)


# ---------------------------------------------------------
# Tables
# ---------------------------------------------------------

print("\nTables\n")

print("Number of Tables :", len(doc.tables))

for table_no, table in enumerate(doc.tables, start=1):

    print(f"\nTable {table_no}")

    rows = []

    for row in table.rows:

        row_data = [cell.text for cell in row.cells]

        rows.append(row_data)

        print(row_data)

    if len(rows) > 1:

        df = pd.DataFrame(rows[1:], columns=rows[0])

        print("\nDataFrame")

        print(df)


# ---------------------------------------------------------
# Images
# ---------------------------------------------------------

print("\nImages\n")

image_count = 0

for rel in doc.part.rels.values():

    if "image" in rel.target_ref:

        image_count += 1

        image_bytes = rel.target_part.blob

        image = Image.open(io.BytesIO(image_bytes))

        plt.imshow(image)

        plt.title(f"Image {image_count}")

        plt.axis("off")

        plt.show()

if image_count == 0:

    print("No Images Found")


# =========================================================
# 2. DOCX2TXT
# =========================================================

print("\n")
print("=" * 70)
print("2. DOCX2TXT")
print("=" * 70)

text = docx2txt.process(FILE)

print("\nExtracted Text\n")

print(text[:1500])

print("\nTotal Characters :", len(text))

print("Total Words :", len(text.split()))

# Extract Images

os.makedirs("docx_images", exist_ok=True)

docx2txt.process(FILE, "docx_images")

print("\nImages (if present) extracted to folder : docx_images")


# =========================================================
# 3. MAMMOTH
# =========================================================

print("\n")
print("=" * 70)
print("3. MAMMOTH")
print("=" * 70)

with open(FILE, "rb") as f:

    result = mammoth.convert_to_html(f)

html = result.value

print("\nGenerated HTML\n")

print(html[:3000])

with open("output.html", "w", encoding="utf-8") as f:

    f.write(html)

print("\nHTML saved as output.html")

if result.messages:

    print("\nMessages")

    for msg in result.messages:

        print(msg)

print("\nDone.")